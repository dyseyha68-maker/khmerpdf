import os
import uuid
import zipfile
import subprocess
import logging
import json
import numpy as np
from celery import shared_task
from django.conf import settings
import fitz
from PIL import Image
import io
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

logger = logging.getLogger(__name__)


def _pdf_has_gs_incompatible_images(path):
    """
    Return True if the PDF contains images that Ghostscript cannot compress
    correctly — specifically, FlateDecode streams that embed PNG data.
    GS mis-interprets the embedded PNG bytes as raw pixels, producing
    blank or corrupted output.  PyMuPDF handles these correctly.
    """
    try:
        doc = fitz.open(path)
        for page in doc:
            for img_info in page.get_images():
                xref = img_info[0]
                xobj = doc.xref_object(xref)
                if '/FlateDecode' in xobj:
                    base = doc.extract_image(xref)
                    if base.get('ext') == 'png':
                        doc.close()
                        return True
        doc.close()
    except Exception:
        pass
    return False


def compress_with_ghostscript(input_path, output_path, compression_level='recommended'):
    """
    Professional PDF compression using Ghostscript with ilovepdf-grade parameters.

    Unlike blunt preset-only commands, we override DPI and JPEG quality independently
    via `-c setdistillerparams`.  This is the key reason ilovepdf "Extreme" still
    looks crisp: they use 96 DPI (not 72) and a moderate JPEG QFactor (not max).

    QFactor reference (Ghostscript / Distiller):
        0.10 → ~q95  (near-lossless)
        0.15 → ~q85  (high quality)
        0.40 → ~q65  (good balance)
        0.76 → ~q45  (ilovepdf "extreme" — readable images, tiny files)
        1.30 → ~q15  (very low quality)
    """

    original_size    = os.path.getsize(input_path)
    original_size_mb = original_size / (1024 * 1024)
    logger.info(f'Compressing {original_size_mb:.2f} MB  level={compression_level}')

    # ── Per-level tuning ───────────────────────────────────────────────────────
    configs = {
        # "Extreme" card — 96 DPI (vs old 72), still big savings
        'low': {
            'pdfsettings': '/screen',
            'color_dpi':    96,
            'gray_dpi':     96,
            'mono_dpi':    150,
        },
        # "Recommended" card
        'recommended': {
            'pdfsettings': '/ebook',
            'color_dpi':   150,
            'gray_dpi':    150,
            'mono_dpi':    200,
        },
        # "Less" card
        'high': {
            'pdfsettings': '/printer',
            'color_dpi':   200,
            'gray_dpi':    200,
            'mono_dpi':    300,
        },
    }

    cfg = configs.get(compression_level, configs['recommended'])

    timeout = max(120, int(original_size_mb * 6))
    timeout = min(timeout, 900)

    # On Windows, Ghostscript is gswin64c / gswin32c; on Linux/Mac it's gs
    import sys as _sys
    if _sys.platform == 'win32':
        import shutil as _shutil
        gs_exe = (
            _shutil.which('gswin64c') or
            _shutil.which('gswin32c') or
            _shutil.which('gs') or
            'gswin64c'
        )
    else:
        gs_exe = 'gs'

    cmd = [
        gs_exe,
        '-sDEVICE=pdfwrite',
        '-dCompatibilityLevel=1.4',
        # ── NO -dPDFSETTINGS preset ───────────────────────────────────────────
        # Presets (/screen, /ebook, etc.) force colour-space conversion to sRGB
        # which turns DeviceCMYK / ICCBased images BLACK on Asian scanned docs.
        # We set only what we need manually so colours are never touched.
        '-dNOPAUSE', '-dQUIET', '-dBATCH', '-dSAFER',

        # ── Image downsampling ────────────────────────────────────────────────
        '-dDownsampleColorImages=true',
        '-dDownsampleGrayImages=true',
        '-dDownsampleMonoImages=true',
        '-dColorImageDownsampleType=/Bicubic',
        '-dGrayImageDownsampleType=/Bicubic',
        '-dMonoImageDownsampleType=/Subsample',
        f'-dColorImageResolution={cfg["color_dpi"]}',
        f'-dGrayImageResolution={cfg["gray_dpi"]}',
        f'-dMonoImageResolution={cfg["mono_dpi"]}',
        '-dColorImageDownsampleThreshold=1.0',
        '-dGrayImageDownsampleThreshold=1.0',
        '-dMonoImageDownsampleThreshold=1.0',

        # ── Structural optimisations ──────────────────────────────────────────
        '-dCompressFonts=true',
        '-dSubsetFonts=true',
        '-dDetectDuplicateImages=true',
        '-dRemoveUnusedResources=true',
        '-dMaxBitmap=500000000',

        # ── Output ────────────────────────────────────────────────────────────
        f'-sOutputFile={output_path}',
        input_path,
    ]

    logger.info(f'GS: dpi={cfg["color_dpi"]}  (no preset — colour-safe)')

    try:
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        try:
            stdout, stderr = proc.communicate(timeout=timeout)
        except subprocess.TimeoutExpired:
            proc.kill()
            proc.communicate()
            raise Exception(f'Compression timed out ({original_size_mb:.1f} MB file)')
        if proc.returncode != 0:
            raise Exception(f'Ghostscript error: {stderr.decode(errors="replace")[:300]}')
    except Exception:
        raise

    if not os.path.exists(output_path):
        raise Exception('Ghostscript produced no output file')

    out_mb = os.path.getsize(output_path) / (1024 * 1024)
    logger.info(f'GS result: {out_mb:.2f} MB (was {original_size_mb:.2f} MB)')

    # If GS made the file larger, fall back
    if os.path.getsize(output_path) >= original_size:
        os.remove(output_path)
        raise Exception('Ghostscript did not reduce file size; trying fallback')

    # ── Blank / black-page guard ──────────────────────────────────────────────
    # GS can produce blank-white or solid-black pages on unusual PDFs.
    # We check output pages — but first check if the SOURCE also renders
    # abnormally via fitz (rotation/encoding quirks). If the source already
    # looks bad in fitz, trust GS; only fall back if source looked normal.
    try:
        def _fitz_page_stats(path, page_num=0):
            """Return (white_ratio, black_ratio) for a page rendered at 36 DPI."""
            try:
                d = fitz.open(path)
                p = d[page_num]
                pix = p.get_pixmap(dpi=36, alpha=False)
                s, n = pix.samples, len(pix.samples) // 3
                d.close()
                if n == 0:
                    return 0.0, 0.0
                w = sum(1 for i in range(0,len(s),3) if s[i]>230 and s[i+1]>230 and s[i+2]>230)
                b = sum(1 for i in range(0,len(s),3) if s[i]<25  and s[i+1]<25  and s[i+2]<25)
                return w/n, b/n
            except Exception:
                return 0.0, 0.0

        src_white, src_black = _fitz_page_stats(input_path)
        src_bad = src_white >= 0.95 or src_black >= 0.90
        if src_bad:
            logger.info(
                f'Source renders abnormally in fitz (white={src_white:.0%} '
                f'black={src_black:.0%}) — trusting GS output'
            )
        else:
            # Source looks normal; verify GS output is not blank/black
            check_doc = fitz.open(output_path)
            bad_count = 0
            total_checked = min(len(check_doc), 3)
            for page in check_doc:
                pix = page.get_pixmap(dpi=36, alpha=False)
                s, n = pix.samples, len(pix.samples) // 3
                if n == 0:
                    bad_count += 1
                else:
                    white_px = sum(1 for i in range(0,len(s),3) if s[i]>230 and s[i+1]>230 and s[i+2]>230)
                    black_px = sum(1 for i in range(0,len(s),3) if s[i]<25  and s[i+1]<25  and s[i+2]<25)
                    if (white_px/n) >= 0.95 or (black_px/n) >= 0.90:
                        bad_count += 1
                        logger.warning(f'GS bad page {page.number}: white={white_px/n:.0%} black={black_px/n:.0%}')
                if page.number + 1 >= total_checked:
                    break
            check_doc.close()
            if bad_count >= total_checked:
                logger.warning('GS produced blank/black pages — falling back to PyMuPDF')
                os.remove(output_path)
                raise Exception('Ghostscript produced blank/black output; trying fallback')
    except fitz.FileNotFoundError:
        raise Exception('Could not verify GS output; trying fallback')

    return output_path


def compress_with_pymupdf(input_path, output_path, compression_level='recommended'):
    """
    Fallback compressor using PyMuPDF.

    Strategy: extract every raster image from the PDF, re-encode it as JPEG at
    the chosen quality/DPI, then replace the stream in-place.  Vectors and text
    are untouched, so readability is preserved.
    """

    # Per-level settings: (jpeg_quality, max_dimension_px)
    # These match the Ghostscript QFactor choices above.
    level_cfg = {
        'low':         (35, 800),    # Extreme  — readable, small
        'recommended': (65, 1600),   # Balanced
        'high':        (82, 2400),   # Less compression
    }
    quality, max_dim = level_cfg.get(compression_level, (65, 1600))

    doc = fitz.open(input_path)
    replaced = 0
    seen_xrefs = set()

    for page in doc:
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                base_img  = doc.extract_image(xref)
                img_bytes = base_img['image']

                # Skip masks / tiny images — not worth re-encoding
                if len(img_bytes) < 2048:
                    continue

                pil_img = Image.open(io.BytesIO(img_bytes))

                # Normalise colour mode
                if pil_img.mode == 'CMYK':
                    pil_img = pil_img.convert('RGB')
                elif pil_img.mode in ('RGBA', 'LA', 'P'):
                    bg = Image.new('RGB', pil_img.size, (255, 255, 255))
                    if pil_img.mode == 'P':
                        pil_img = pil_img.convert('RGBA')
                    if pil_img.mode in ('RGBA', 'LA'):
                        bg.paste(pil_img, mask=pil_img.split()[-1])
                    else:
                        bg.paste(pil_img)
                    pil_img = bg
                elif pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')

                # ── Inverted-image guard ──────────────────────────────────
                # Some PDFs (e.g. rotated scanner outputs) store images as
                # photographic negatives. Mean < 20 → almost all pixels are
                # near-black → invert so the output is readable.
                arr = np.array(pil_img)
                if arr.mean() < 20:
                    pil_img = Image.fromarray(255 - arr)
                    logger.info(f'Inverted near-black image xref={xref}')

                # Downsample if larger than max_dim
                if pil_img.width > max_dim or pil_img.height > max_dim:
                    pil_img.thumbnail((max_dim, max_dim), Image.LANCZOS)

                new_w, new_h = pil_img.size

                # Re-encode as JPEG
                buf = io.BytesIO()
                pil_img.save(buf, format='JPEG', quality=quality,
                             optimize=True, progressive=True)
                new_bytes = buf.getvalue()

                # Only replace if we actually saved space
                if len(new_bytes) >= len(img_bytes):
                    continue

                # ── Proper JPEG injection ──────────────────────────────────
                # update_stream(compress=True) wraps JPEG in FlateDecode and
                # leaves the original filter untouched → black/corrupted pages.
                # Correct approach:
                #   1. Store raw JPEG bytes (compress=False)
                #   2. Patch the XObject dict to reflect DCTDecode + new dims
                doc.update_stream(xref, new_bytes, compress=False)
                doc.xref_set_key(xref, 'Filter', '/DCTDecode')
                doc.xref_set_key(xref, 'Width',  str(new_w))
                doc.xref_set_key(xref, 'Height', str(new_h))
                try:
                    if 'DecodeParms' in doc.xref_object(xref):
                        doc.xref_del_key(xref, 'DecodeParms')
                except Exception:
                    pass
                replaced += 1

            except Exception:
                continue   # skip unprocessable image, leave original

    logger.info(f'PyMuPDF: replaced {replaced} image streams')

    # Save with maximum structural compression
    try:
        doc.save(output_path, deflate=True, deflate_images=True, garbage=4, clean=True)
    except Exception:
        doc.save(output_path, deflate=True, garbage=4, clean=True)
    doc.close()


@shared_task
def compress_pdf(job_id):
    from apps.pdf.models import Job
    from django.db import close_old_connections
    close_old_connections()   # ensure fresh DB connection in this thread

    Job.objects.filter(id=job_id).update(status='processing')
    job = Job.objects.get(id=job_id)
    
    compression_level = job.compression_level or 'recommended'
    
    try:
        file_ids = job.files if job.files else []
        
        if file_ids and len(file_ids) > 1:
            # Multiple files - create ZIP
            first_base_name = ''
            for idx, file_id in enumerate(file_ids):
                try:
                    job_file = Job.objects.get(id=file_id)
                    if job_file.file and os.path.exists(job_file.file.path):
                        if idx == 0:
                            first_base_name = os.path.splitext(os.path.basename(job_file.file.name))[0]
                        break
                except:
                    continue
            
            zip_filename = f'{first_base_name}_compressed.zip' if first_base_name else f'compressed_{uuid.uuid4().hex[:8]}.zip'
            zip_path = os.path.join(settings.MEDIA_ROOT, 'processed', zip_filename)
            os.makedirs(os.path.dirname(zip_path), exist_ok=True)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for file_id in file_ids:
                    try:
                        job_file = Job.objects.get(id=file_id)
                        if job_file.file and os.path.exists(job_file.file.path):
                            input_path = job_file.file.path
                            original_size = os.path.getsize(input_path)
                            
                            base_name = os.path.splitext(os.path.basename(job_file.file.name))[0]
                            output_filename = f'{base_name}_compressed.pdf'
                            output_path = os.path.join(settings.MEDIA_ROOT, 'processed', f'{uuid.uuid4().hex[:8]}_compressed.pdf')
                            
                            # Skip GS for PDFs with FlateDecode+PNG images
                            if _pdf_has_gs_incompatible_images(input_path):
                                logger.info(f'FlateDecode+PNG — using PyMuPDF directly for {base_name}')
                                compress_with_pymupdf(input_path, output_path, compression_level)
                            else:
                                try:
                                    compress_with_ghostscript(input_path, output_path, compression_level)
                                    logger.info(f'Used Ghostscript for {base_name}')
                                except Exception as gs_err:
                                    logger.warning(f'Ghostscript failed, falling back to PyMuPDF: {gs_err}')
                                    compress_with_pymupdf(input_path, output_path, compression_level)
                            
                            compressed_size = os.path.getsize(output_path)
                            if compressed_size >= original_size:
                                # Use original if compression made it larger
                                logger.info(f'Compression did not help for {base_name}; using original')
                                os.remove(output_path)
                                zip_file.write(input_path, output_filename)
                            else:
                                zip_file.write(output_path, output_filename)
                                os.remove(output_path)
                    except Exception as e:
                        continue
            
            job.result.save(zip_filename, open(zip_path, 'rb'))
            os.remove(zip_path)
            Job.objects.filter(id=job_id).update(status='done')
            return {'status': 'done', 'job_id': str(job_id)}
        
        else:
            # Single file
            input_path = job.file.path
            original_size = os.path.getsize(input_path)
            base_name = os.path.splitext(os.path.basename(job.file.name))[0]
            output_filename = f'{base_name}_compressed.pdf'
            output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Skip GS for PDFs with FlateDecode+PNG images (GS produces blank output)
            if _pdf_has_gs_incompatible_images(input_path):
                logger.info('FlateDecode+PNG detected — using PyMuPDF directly')
                try:
                    compress_with_pymupdf(input_path, output_path, compression_level)
                except Exception as pymupdf_err:
                    logger.error(f'PyMuPDF failed: {pymupdf_err}')
                    raise Exception(f'Compression failed: {pymupdf_err}')
            else:
                try:
                    compress_with_ghostscript(input_path, output_path, compression_level)
                    logger.info('Used Ghostscript for single file')
                except Exception as gs_err:
                    logger.warning(f'Ghostscript failed, falling back to PyMuPDF: {gs_err}')
                    try:
                        compress_with_pymupdf(input_path, output_path, compression_level)
                    except Exception as pymupdf_err:
                        logger.error(f'Both compression methods failed: GS={gs_err}, PyMuPDF={pymupdf_err}')
                        raise Exception(f'Compression failed: {pymupdf_err}')
            
            if not os.path.exists(output_path):
                raise Exception(f'Compression failed: output file not created')

            compressed_size = os.path.getsize(output_path)

            if compressed_size >= original_size:
                # Compression made it larger or equal — return original unchanged
                logger.info(
                    f'Compression did not reduce size ({compressed_size} >= {original_size}); '
                    'returning original file'
                )
                os.remove(output_path)
                job.result.save(output_filename, open(input_path, 'rb'))
                compressed_size = original_size
            else:
                job.result.save(output_filename, open(output_path, 'rb'))
                os.remove(output_path)

            Job.objects.filter(id=job_id).update(status='done')

            return {'status': 'done', 'job_id': str(job_id), 'original': original_size, 'compressed': compressed_size}

    except Exception as e:
        logger.error(f'compress_pdf failed: {e}', exc_info=True)
        Job.objects.filter(id=job_id).update(status='failed', error_message=str(e)[:500])
        raise


def merge_pdf(job_id):
    from apps.pdf.models import Job
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()
    
    try:
        file_paths = job.files if job.files else []
        
        first_base_name = ''
        for idx, file_id in enumerate(file_paths):
            try:
                job_file = Job.objects.get(id=file_id)
                if job_file.file and os.path.exists(job_file.file.path):
                    if idx == 0:
                        first_base_name = os.path.splitext(os.path.basename(job_file.file.name))[0]
                    break
            except:
                continue
        
        output_filename = f'{first_base_name}_merged.pdf' if first_base_name else f'merged_{uuid.uuid4().hex[:8]}.pdf'
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        
        if job.result:
            job.result.delete()
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        writer = PdfWriter()
        
        for file_id in file_paths:
            try:
                job_file = Job.objects.get(id=file_id)
                if job_file.file and os.path.exists(job_file.file.path):
                    reader = PdfReader(job_file.file.path)
                    for page in reader.pages:
                        writer.add_page(page)
            except:
                continue
        
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        job.result.save(output_filename, open(output_path, 'rb'))
        
        os.remove(output_path)
        
        job.status = 'done'
        job.save()
        
        return {'status': 'done', 'job_id': str(job_id)}
        
    except Exception as e:
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise


def split_pdf(job_id):
    from apps.pdf.models import Job
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()
    
    try:
        input_path = job.file.path
        page_range = job.page_range or '1'
        split_mode = job.compression_level or 'range'
        
        base_name = os.path.splitext(os.path.basename(job.file.name))[0]
        
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        if split_mode == 'every':
            import zipfile
            n = 3
            if page_range.startswith('every:'):
                n = int(page_range.split(':')[1]) or 3
            
            zip_filename = f'{base_name}_split.zip'
            zip_path = os.path.join(settings.MEDIA_ROOT, 'processed', zip_filename)
            os.makedirs(os.path.dirname(zip_path), exist_ok=True)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for i in range(0, total_pages, n):
                    writer = PdfWriter()
                    for j in range(i, min(i + n, total_pages)):
                        writer.add_page(reader.pages[j])
                    
                    page_path = os.path.join(settings.MEDIA_ROOT, 'processed', f'page_{i//n+1}.pdf')
                    with open(page_path, 'wb') as f:
                        writer.write(f)
                    
                    zip_file.write(page_path, f'pages_{i+1}-{min(i+n,total_pages)}.pdf')
                    os.remove(page_path)
            
            job.result.save(zip_filename, open(zip_path, 'rb'))
            os.remove(zip_path)
            
        else:
            output_filename = f'{base_name}_split.pdf'
            output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            writer = PdfWriter()
            pages = parse_page_range(page_range, total_pages)
            
            for page_num in pages:
                if 0 <= page_num < total_pages:
                    writer.add_page(reader.pages[page_num])
            
            with open(output_path, 'wb') as f:
                writer.write(f)
            
            job.result.save(output_filename, open(output_path, 'rb'))
            os.remove(output_path)
        
        job.status = 'done'
        job.save()
        
        return {'status': 'done', 'job_id': str(job_id)}
        
    except Exception as e:
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise


@shared_task
def protect_pdf(job_id, password):
    """Lock a PDF (or batch of PDFs) with an open password using pikepdf
    (qpdf-backed), AES-256 (R=6 — the modern PDF 2.0 encryption revision).
    `password` arrives only as a task argument — it is never written to
    the Job row, so it never lands in the database or the admin panel."""
    from apps.pdf.models import Job
    from django.db import close_old_connections
    close_old_connections()
    import pikepdf

    Job.objects.filter(id=job_id).update(status='processing')
    job = Job.objects.get(id=job_id)

    def _encrypt_one(input_path, output_path):
        try:
            with pikepdf.open(input_path) as pdf:
                pdf.save(
                    output_path,
                    encryption=pikepdf.Encryption(owner=password, user=password, R=6),
                )
        except pikepdf.PasswordError:
            raise Exception('This PDF is already password-protected — remove the existing password first.')

    try:
        file_ids = job.files if job.files else []

        if file_ids and len(file_ids) > 1:
            # Multiple files - create ZIP
            first_base_name = ''
            for idx, file_id in enumerate(file_ids):
                try:
                    job_file = Job.objects.get(id=file_id)
                    if job_file.file and os.path.exists(job_file.file.path):
                        if idx == 0:
                            first_base_name = os.path.splitext(os.path.basename(job_file.file.name))[0]
                        break
                except Exception:
                    continue

            zip_filename = f'{first_base_name}_protected.zip' if first_base_name else f'protected_{uuid.uuid4().hex[:8]}.zip'
            zip_path = os.path.join(settings.MEDIA_ROOT, 'processed', zip_filename)
            os.makedirs(os.path.dirname(zip_path), exist_ok=True)

            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for file_id in file_ids:
                    try:
                        job_file = Job.objects.get(id=file_id)
                        if job_file.file and os.path.exists(job_file.file.path):
                            input_path = job_file.file.path
                            base_name = os.path.splitext(os.path.basename(job_file.file.name))[0]
                            output_filename = f'{base_name}_protected.pdf'
                            output_path = os.path.join(settings.MEDIA_ROOT, 'processed', f'{uuid.uuid4().hex[:8]}_protected.pdf')

                            _encrypt_one(input_path, output_path)

                            zip_file.write(output_path, output_filename)
                            os.remove(output_path)
                    except Exception as e:
                        logger.warning(f'protect_pdf: skipping file {file_id}: {e}')
                        continue

            job.result.save(zip_filename, open(zip_path, 'rb'))
            os.remove(zip_path)
            Job.objects.filter(id=job_id).update(status='done')
            return {'status': 'done', 'job_id': str(job_id)}

        else:
            input_path = job.file.path
            base_name = os.path.splitext(os.path.basename(job.file.name))[0]
            output_filename = f'{base_name}_protected.pdf'
            output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            _encrypt_one(input_path, output_path)

            job.result.save(output_filename, open(output_path, 'rb'))
            os.remove(output_path)

            Job.objects.filter(id=job_id).update(status='done')
            return {'status': 'done', 'job_id': str(job_id)}

    except Exception as e:
        logger.error(f'protect_pdf failed: {e}', exc_info=True)
        Job.objects.filter(id=job_id).update(status='failed', error_message=str(e)[:500])
        raise


def parse_page_range(page_str, total_pages):
    pages = set()
    
    parts = page_str.split(',')
    for part in parts:
        part = part.strip()
        if '-' in part:
            start, end = part.split('-')
            start = int(start.strip()) - 1
            end = int(end.strip()) - 1
            for p in range(max(0, start), min(end + 1, total_pages)):
                pages.add(p)
        else:
            p = int(part.strip()) - 1
            if 0 <= p < total_pages:
                pages.add(p)
    
    return sorted(list(pages))


@shared_task
def organize_pdf(job_id, replace_files=None):
    from apps.pdf.models import Job
    import json
    
    if replace_files is None:
        replace_files = {}
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()
    
    try:
        input_path = job.file.path
        page_order_json = job.page_range or '[]'
        
        try:
            page_order = json.loads(page_order_json.replace('\\/', '/'))
        except:
            page_order = []
        
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        output_filename = f'organized_{uuid.uuid4().hex[:8]}.pdf'
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        writer = PdfWriter()
        
        for page_info in page_order:
            page_idx = page_info.get('index')
            is_blank = page_info.get('isBlank', False)
            is_replaced = page_info.get('isReplaced', False)
            replaced_file_index = page_info.get('replacedFileIndex')
            replaced_page_num = page_info.get('replacedPageNum')
            
            if is_blank:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                temp_blank = os.path.join(settings.MEDIA_ROOT, 'processed', f'blank_{uuid.uuid4().hex[:8]}.pdf')
                c = canvas.Canvas(temp_blank, pagesize=letter)
                c.showPage()
                c.save()
                
                blank_reader = PdfReader(temp_blank)
                writer.add_page(blank_reader.pages[0])
                os.remove(temp_blank)
            elif is_replaced and replaced_file_index is not None and replaced_page_num is not None:
                replace_file = replace_files.get(replaced_file_index)
                if replace_file:
                    replace_path = os.path.join(settings.MEDIA_ROOT, 'uploads', replace_file.name)
                    os.makedirs(os.path.dirname(replace_path), exist_ok=True)
                    with open(replace_path, 'wb') as f:
                        for chunk in replace_file.chunks():
                            f.write(chunk)
                    
                    replace_reader = PdfReader(replace_path)
                    if 0 <= replaced_page_num - 1 < len(replace_reader.pages):
                        writer.add_page(replace_reader.pages[replaced_page_num - 1])
                    os.remove(replace_path)
            else:
                orig_page = page_info.get('original')
                if orig_page and 0 <= orig_page - 1 < total_pages:
                    writer.add_page(reader.pages[orig_page - 1])
                elif page_idx is not None and 0 <= page_idx < total_pages:
                    writer.add_page(reader.pages[page_idx])
        
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        job.result.save(output_filename, open(output_path, 'rb'))
        os.remove(output_path)
        
        job.status = 'done'
        job.save()
        
        return {'status': 'done', 'job_id': str(job_id)}
        
    except Exception as e:
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise


# ── Google Cloud Vision OCR (optional, admin-configured) ───────────────────────
# Off by default. When apps.pdf.models.OCRSettings.use_google_vision is turned
# on (with a valid service account JSON pasted into the admin), ocr_pdf below
# uses Vision's DOCUMENT_TEXT_DETECTION instead of Tesseract, per page, with
# an automatic per-page fallback to Tesseract on any failure (network, quota,
# bad credentials) unless the admin has switched that fallback off.

_VISION_LANG_HINTS = {
    'eng': ['en'],
    'khm': ['km'],
    'eng+khm': ['en', 'km'],
}


def _get_ocr_settings():
    from apps.pdf.models import OCRSettings
    return OCRSettings.objects.filter(pk=1).first()


def _build_vision_client(ocr_settings):
    """Build an authenticated Vision client from the admin-pasted service
    account JSON. Raises if the JSON is missing/malformed so the caller can
    log a clear error and fall back to Tesseract."""
    from google.cloud import vision
    from google.oauth2 import service_account

    info = json.loads(ocr_settings.google_credentials_json)
    credentials = service_account.Credentials.from_service_account_info(info)
    return vision.ImageAnnotatorClient(credentials=credentials)


def _vision_ocr_page(client, pil_image, lang_hints=None):
    """Run Google Vision DOCUMENT_TEXT_DETECTION on one page image, return
    the extracted text (same shape as pytesseract.image_to_string output)."""
    from google.cloud import vision

    buf = io.BytesIO()
    pil_image.convert('RGB').save(buf, format='PNG')
    image = vision.Image(content=buf.getvalue())
    image_context = vision.ImageContext(language_hints=lang_hints) if lang_hints else None

    response = client.document_text_detection(image=image, image_context=image_context)
    if response.error and response.error.message:
        raise RuntimeError(response.error.message)
    return response.full_text_annotation.text


@shared_task
def ocr_pdf(job_id):
    from apps.pdf.models import Job
    import logging
    import pytesseract
    from PIL import Image
    logger = logging.getLogger(__name__)
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()

    ocr_lang = job.compression_level or 'eng'

    lang_map = {
        'eng': 'eng',
        'khm': 'khm',
        'eng+khm': 'eng+khm'
    }
    tess_lang = lang_map.get(ocr_lang, 'eng')

    tesseract_cmd = getattr(settings, 'TESSERACT_CMD', None)
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    
    try:
        input_path = job.file.path
        file_size_mb = os.path.getsize(input_path) / (1024 * 1024)
        logger.info(f'Starting OCR to editable Word, file size: {file_size_mb:.1f}MB')
        
        from pdf2image import convert_from_path
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        poppler_path = getattr(settings, 'POPPLER_PATH', None)
        pages = convert_from_path(input_path, dpi=150, poppler_path=poppler_path)
        max_pages = min(len(pages), 30)
        
        def clean_text(text):
            if not text:
                return ""
            lines = text.split('\n')
            cleaned = []
            for line in lines:
                line = line.strip()
                if len(line) < 2:
                    continue
                line = line.rstrip('|_»¿¡')
                if line:
                    cleaned.append(line)
            return '\n'.join(cleaned)
        
        # Set table auto-style
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        extracted_chars = 0

        # ── Optional Google Cloud Vision engine (admin-configured) ──
        ocr_settings = _get_ocr_settings()
        vision_client = None
        vision_fallback_ok = True
        if ocr_settings and ocr_settings.use_google_vision and ocr_settings.google_credentials_json.strip():
            vision_fallback_ok = ocr_settings.fallback_to_tesseract
            try:
                vision_client = _build_vision_client(ocr_settings)
                logger.info('OCR: using Google Cloud Vision (DOCUMENT_TEXT_DETECTION)')
            except Exception as e:
                logger.error(f'OCR: Google Vision client init failed: {e}', exc_info=True)
                if not vision_fallback_ok:
                    raise RuntimeError(f'Google Vision is enabled but could not be configured: {e}')
                vision_client = None  # falls through to Tesseract below
        vision_lang_hints = _VISION_LANG_HINTS.get(ocr_lang, ['en'])

        for i in range(max_pages):
            logger.info(f'Processing page {i+1}/{max_pages}')
            page = pages[i]
            gray = page.convert('L')
            page_done = False

            # ── Try Google Vision first, if configured ──
            if vision_client:
                try:
                    vtext = clean_text(_vision_ocr_page(vision_client, page, lang_hints=vision_lang_hints))
                    if vtext:
                        for line in vtext.split('\n'):
                            line = line.strip()
                            if not line:
                                continue
                            p = doc.add_paragraph(line)
                            extracted_chars += len(line)
                            has_khmer = any('ក' <= c <= '៿' for c in line)
                            for run in p.runs:
                                run.font.name = 'Kantumruy Pro' if has_khmer else 'Calibri'
                                run.font.size = Pt(11)
                    page_done = True
                except Exception as e:
                    logger.error(f'Google Vision OCR failed on page {i+1}: {e}', exc_info=True)
                    if not vision_fallback_ok:
                        raise RuntimeError(f'Google Vision OCR failed on page {i+1}: {e}')
                    # else: fall through to Tesseract for this page

            # Get OCR with position data
            data = {'left': [], 'top': [], 'width': [], 'height': [], 'text': []}
            if not page_done:
                try:
                    data = pytesseract.image_to_data(gray, lang=tess_lang, output_type=pytesseract.Output.DICT)
                except Exception as e:
                    logger.error(f'pytesseract.image_to_data failed on page {i+1}: {e}', exc_info=True)
                    data = {'left': [], 'top': [], 'width': [], 'height': [], 'text': []}

            # Extract text and organize by lines (y-position)
            lines_dict = {}
            n = 0 if page_done else len(data.get('text', []))

            if n > 0:
                # Group words into lines based on y-position
                for j in range(n):
                    text = data['text'][j].strip()
                    if not text:
                        continue
                    y = int(data['top'][j])
                    conf = int(data['conf'][j])
                    if conf < 30:
                        continue
                    
                    # Group to nearest line (within 10 pixels)
                    found_line = False
                    for line_y in lines_dict:
                        if abs(y - line_y) < 10:
                            lines_dict[line_y].append(data['text'][j])
                            found_line = True
                            break
                    if not found_line:
                        lines_dict[y] = [data['text'][j]]
                
                # Sort lines by y-position and add to document
                for line_y in sorted(lines_dict.keys()):
                    line_text = ' '.join(lines_dict[line_y])
                    line_text = line_text.strip()
                    if line_text:
                        p = doc.add_paragraph(line_text)
                        extracted_chars += len(line_text)

                        # Detect Khmer
                        has_khmer = any('\u1780' <= c <= '\u17FF' for c in line_text)

                        # Set font
                        for run in p.runs:
                            if has_khmer:
                                run.font.name = 'Kantumruy Pro'
                            else:
                                run.font.name = 'Calibri'
                            run.font.size = Pt(11)
            elif not page_done:
                # Fallback to basic OCR (Tesseract only — Vision already had
                # its turn above if it was configured, page_done covers that)
                try:
                    text = pytesseract.image_to_string(gray, lang=tess_lang, config='--psm 6')
                    text = clean_text(text)
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                p = doc.add_paragraph(line)
                                extracted_chars += len(line.strip())
                                for run in p.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(11)
                except Exception as e:
                    logger.error(f'pytesseract.image_to_string failed on page {i+1}: {e}', exc_info=True)

            if i < max_pages - 1:
                doc.add_page_break()

        if extracted_chars == 0:
            raise RuntimeError(
                'OCR produced no text on any page. This usually means Tesseract-OCR '
                'is not installed / not found (check TESSERACT_CMD in settings.py or '
                'that tesseract is on PATH), the wrong language pack is missing '
                f"(requested '{tess_lang}'), or the source pages are blank/too low quality."
            )

        output_filename = f'ocr_{uuid.uuid4().hex[:8]}.docx'
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)

        job.result.save(output_filename, open(output_path, 'rb'))
        os.remove(output_path)

        job.status = 'done'
        job.save()

        return {'status': 'done', 'job_id': str(job_id)}

    except Exception as e:
        logger.error(f'OCR error: {e}', exc_info=True)
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise


@shared_task
def pdf_to_image_task(job_id):
    from apps.pdf.models import Job
    import zipfile
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()
    
    try:
        input_path = job.file.path
        image_format = job.page_range or 'png'
        dpi = int(job.compression_level or '300')
        
        from pdf2image import convert_from_path

        logger.info(f'Converting PDF to images, format: {image_format}, DPI: {dpi}')

        poppler_path = getattr(settings, 'POPPLER_PATH', None)
        pages = convert_from_path(input_path, dpi=dpi, poppler_path=poppler_path)

        base_name = os.path.splitext(os.path.basename(job.file.name))[0]
        zip_filename = f'{base_name}_images.zip'
        zip_path = os.path.join(settings.MEDIA_ROOT, 'processed', zip_filename)
        os.makedirs(os.path.dirname(zip_path), exist_ok=True)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, page in enumerate(pages):
                img_bytes = io.BytesIO()
                ext = 'jpg' if image_format == 'jpg' else 'png'
                page.save(img_bytes, format=ext.upper())
                img_bytes.seek(0)
                zip_file.writestr(f'page_{i+1:03d}.{ext}', img_bytes.read())
        
        job.result.save(zip_filename, open(zip_path, 'rb'))
        os.remove(zip_path)
        
        job.status = 'done'
        job.save()
        
        return {'status': 'done', 'job_id': str(job_id)}
        
    except Exception as e:
        logger.error(f'PDF to Image error: {e}', exc_info=True)
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise


@shared_task
def image_to_pdf_task(file_paths, job_id):
    from apps.pdf.models import Job
    
    job = Job.objects.get(id=job_id)
    job.status = 'processing'
    job.save()
    
    try:
        first_base_name = ''
        if file_paths:
            first_base_name = os.path.splitext(os.path.basename(file_paths[0]))[0]
        
        output_filename = f'{first_base_name}.pdf' if first_base_name else f'images_{uuid.uuid4().hex[:8]}.pdf'
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        images = []
        for img_path in file_paths:
            try:
                if os.path.exists(img_path):
                    img = Image.open(img_path)
                    if img.mode == 'RGBA':
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[3])
                        img = background
                    elif img.mode != 'RGB':
                        img = img.convert('RGB')
                    images.append(img)
            except Exception as e:
                logger.error(f'Error loading image {img_path}: {e}')
                continue
        
        if images:
            images[0].save(output_path, save_all=True, append_images=images[1:])
            job.result.save(output_filename, open(output_path, 'rb'))
            os.remove(output_path)
        
        job.status = 'done'
        job.save()
        
        return {'status': 'done', 'job_id': str(job_id)}
        
    except Exception as e:
        logger.error(f'Image to PDF error: {e}', exc_info=True)
        job.status = 'failed'
        job.error_message = str(e)
        job.save()
        raise
